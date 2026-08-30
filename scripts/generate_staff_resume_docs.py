from __future__ import annotations

import argparse
import json
import os
import re
from datetime import date, datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
import pymysql


DOC_FONT = "STHeiti"
_DATABASE_PATTERN = re.compile(r"lu_test_[a-z0-9_]+")


def _require_validation_database() -> str:
    database = os.getenv("DB_DATABASE", "").strip()
    if not _DATABASE_PATTERN.fullmatch(database):
        raise ValueError("staff resume generation requires a lu_test_* database")
    if os.getenv("APP_ENV", "development").strip().lower() in {"prod", "production"}:
        raise ValueError("staff resume generation requires a development validation profile")
    return database


def _apply_font(run, size: Pt | None = None) -> None:
    run.font.name = DOC_FONT
    run._element.rPr.rFonts.set(qn("w:ascii"), DOC_FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), DOC_FONT)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), DOC_FONT)
    if size is not None:
        run.font.size = size


def _connect() -> pymysql.connections.Connection:
    database = _require_validation_database()
    return pymysql.connect(
        host=os.getenv("DB_HOST", "127.0.0.1"),
        port=int(os.getenv("DB_PORT", "3306")),
        user=os.getenv("DB_USER", "root"),
        password=os.getenv("DB_PASSWORD", "1234"),
        database=database,
        charset="utf8mb4",
        cursorclass=pymysql.cursors.DictCursor,
    )


def _age(value: date | datetime | str | None) -> int | None:
    if not value:
        return None
    if isinstance(value, str):
        value = datetime.strptime(value, "%Y-%m-%d").date()
    if isinstance(value, datetime):
        value = value.date()
    today = date.today()
    return today.year - value.year - ((today.month, today.day) < (value.month, value.day))


def _join(values: list[str] | None) -> str:
    values = [str(v).strip() for v in (values or []) if str(v).strip()]
    return "、".join(values) if values else "未填"


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _set_cell_text(cell, text: str, *, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    _apply_font(run, Pt(10))
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _add_kv_table(doc: Document, rows: list[tuple[str, str]], widths: tuple[float, float] = (3.2, 12.4)) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    table.autofit = False
    for row_idx, (label, value) in enumerate(rows):
        cells = table.rows[row_idx].cells
        cells[0].width = Cm(widths[0])
        cells[1].width = Cm(widths[1])
        _set_cell_shading(cells[0], "EAF2F8")
        _set_cell_text(cells[0], label, bold=True, color="1F4E79")
        _set_cell_text(cells[1], value)
    doc.add_paragraph()


def _add_heading(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    _apply_font(run, Pt(13))
    run.font.color.rgb = RGBColor(31, 78, 121)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(4)


def _fetch_staff(limit: int) -> list[dict]:
    import pymysql

    conn = _connect()
    try:
        with conn.cursor() as cur:
            cur.execute(
                """
                SELECT s.id,s.name,s.phone,s.email,s.birthday,s.city,s.zip_code,s.address,
                       s.has_massage_cert,s.status,s.weekly_rest_days,s.care_babies,
                       s.service_regions,s.special_skills,
                       COUNT(DISTINCT ss.work_date) AS scheduled_days,
                       COUNT(DISTINCT csa.case_no) AS assigned_cases
                FROM staff s
                LEFT JOIN staff_schedule ss ON ss.staff_id=s.id AND ss.is_work_day=1
                LEFT JOIN case_staff_assignments csa ON csa.staff_id=s.id
                GROUP BY s.id
                ORDER BY s.id
                LIMIT %s
                """,
                (limit,),
            )
            rows = list(cur.fetchall())
            ids = [row["id"] for row in rows]

            def group(table: str, column: str) -> dict[int, list[str]]:
                if not ids:
                    return {}
                placeholders = ",".join(["%s"] * len(ids))
                cur.execute(
                    f"SELECT staff_id,{column} AS val FROM {table} "
                    f"WHERE staff_id IN ({placeholders}) ORDER BY staff_id,{column}",
                    ids,
                )
                out: dict[int, list[str]] = {}
                for item in cur.fetchall():
                    out.setdefault(int(item["staff_id"]), []).append(str(item["val"]))
                return out

            extras = {
                "regions": group("staff_regions", "region_name"),
                "slots": group("staff_time_slots", "slot_name"),
                "skills": group("staff_cooking_skills", "skill_name"),
                "transport": group("staff_transportation", "vehicle_type"),
                "babies": group("staff_baby_types", "baby_type"),
                "rests": group("staff_weekly_rest", "rest_type"),
            }
            for row in rows:
                for key, values in extras.items():
                    row[key] = values.get(row["id"], [])
            return rows
    finally:
        conn.close()


def _resume_doc(staff: dict) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    normal = doc.styles["Normal"]
    normal.font.name = DOC_FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), DOC_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), DOC_FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), DOC_FONT)
    normal.font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"月嫂履歷資歷表｜{staff['name']}")
    run.bold = True
    _apply_font(run, Pt(20))
    run.font.color.rgb = RGBColor(31, 78, 121)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = subtitle.add_run(f"內部派案測試履歷｜Staff ID {staff['id']}｜假資料")
    _apply_font(sub, Pt(10))
    sub.font.color.rgb = RGBColor(90, 90, 90)

    _add_heading(doc, "一、基本資料")
    birthday = staff.get("birthday")
    birthday_text = str(birthday) if birthday else "未填"
    staff_age = _age(birthday)
    if staff_age is not None:
        birthday_text += f"（約 {staff_age} 歲）"
    _add_kv_table(
        doc,
        [
            ("姓名", staff["name"]),
            ("聯絡電話", staff.get("phone") or "未填"),
            ("電子郵件", staff.get("email") or "未填"),
            ("居住地", f"{staff.get('city') or ''} {staff.get('zip_code') or ''} {staff.get('address') or ''}".strip()),
            ("狀態", "可派案" if staff.get("status") == "active" else str(staff.get("status") or "未填")),
            ("按摩證照", "具備" if staff.get("has_massage_cert") else "未填/未具備"),
        ],
    )

    _add_heading(doc, "二、服務能力與偏好")
    _add_kv_table(
        doc,
        [
            ("服務區域", _join(staff.get("regions"))),
            ("可服務時段", _join(staff.get("slots"))),
            ("照護寶寶類型", _join(staff.get("babies"))),
            ("料理/專長", _join(staff.get("skills"))),
            ("交通方式", _join(staff.get("transport"))),
            ("休假偏好", _join(staff.get("rests"))),
        ],
    )

    _add_heading(doc, "三、派案資歷摘要")
    scheduled = int(staff.get("scheduled_days") or 0)
    cases = int(staff.get("assigned_cases") or 0)
    profile = [
        f"目前系統紀錄已排班工作日數：{scheduled} 日。",
        f"目前系統紀錄已關聯案件數：{cases} 件。",
        f"可照護寶寶數上限：{staff.get('care_babies') or '未填'}。",
    ]
    for item in profile:
        p = doc.add_paragraph(style=None)
        p.style = doc.styles["Normal"]
        p.paragraph_format.left_indent = Cm(0.4)
        p.paragraph_format.space_after = Pt(3)
        _apply_font(p.add_run(item), Pt(10.5))

    _add_heading(doc, "四、內部派案備註")
    note = doc.add_paragraph()
    _apply_font(note.add_run(
        "本履歷由目前本機測試資料庫自動生成，供 UI/API 驗收、派案流程展示與文件版型測試使用。"
        "正式對外使用前，需由行政人員重新核對本人資料、證照、服務經驗與可派案時段。"
    ), Pt(10.5))

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run(f"新竹市月子工會內部測試文件｜生成日 {date.today().isoformat()}")
    _apply_font(footer_run, Pt(8))
    footer_run.font.color.rgb = RGBColor(120, 120, 120)
    return doc


def _resume_pdf(staff: dict, path: Path) -> None:
    pdf_font = "ResumeCJK"
    font_path = "/System/Library/Fonts/STHeiti Medium.ttc"
    if pdf_font not in pdfmetrics.getRegisteredFontNames():
        pdfmetrics.registerFont(TTFont(pdf_font, font_path))
    styles = getSampleStyleSheet()
    base = ParagraphStyle(
        "CJKBase",
        parent=styles["Normal"],
        fontName=pdf_font,
        fontSize=10.5,
        leading=15,
        textColor=colors.HexColor("#222222"),
    )
    title_style = ParagraphStyle(
        "CJKTitle",
        parent=base,
        fontSize=20,
        leading=26,
        alignment=1,
        textColor=colors.HexColor("#1F4E79"),
        spaceAfter=4,
    )
    subtitle_style = ParagraphStyle(
        "CJKSubtitle",
        parent=base,
        fontSize=10,
        leading=14,
        alignment=1,
        textColor=colors.HexColor("#666666"),
        spaceAfter=14,
    )
    heading_style = ParagraphStyle(
        "CJKHeading",
        parent=base,
        fontSize=13,
        leading=18,
        textColor=colors.HexColor("#1F4E79"),
        spaceBefore=8,
        spaceAfter=5,
    )

    def p(text: str, style: ParagraphStyle = base) -> Paragraph:
        return Paragraph(str(text).replace("\n", "<br/>"), style)

    def kv(rows: list[tuple[str, str]]) -> Table:
        table = Table(
            [[p(label), p(value)] for label, value in rows],
            colWidths=[3.6 * cm, 11.8 * cm],
            hAlign="LEFT",
        )
        table.setStyle(
            TableStyle(
                [
                    ("FONTNAME", (0, 0), (-1, -1), pdf_font),
                    ("FONTSIZE", (0, 0), (-1, -1), 10),
                    ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#EAF2F8")),
                    ("TEXTCOLOR", (0, 0), (0, -1), colors.HexColor("#1F4E79")),
                    ("GRID", (0, 0), (-1, -1), 0.45, colors.HexColor("#666666")),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 7),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 7),
                    ("TOPPADDING", (0, 0), (-1, -1), 5),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ]
            )
        )
        return table

    birthday = staff.get("birthday")
    birthday_text = str(birthday) if birthday else "未填"
    staff_age = _age(birthday)
    if staff_age is not None:
        birthday_text += f"（約 {staff_age} 歲）"

    story = [
        p(f"月嫂履歷資歷表｜{staff['name']}", title_style),
        p(f"內部派案測試履歷｜Staff ID {staff['id']}｜假資料", subtitle_style),
        p("一、基本資料", heading_style),
        kv(
            [
                ("姓名", staff["name"]),
                ("聯絡電話", staff.get("phone") or "未填"),
                ("電子郵件", staff.get("email") or "未填"),
                ("生日", birthday_text),
                ("居住地", f"{staff.get('city') or ''} {staff.get('zip_code') or ''} {staff.get('address') or ''}".strip()),
                ("狀態", "可派案" if staff.get("status") == "active" else str(staff.get("status") or "未填")),
                ("按摩證照", "具備" if staff.get("has_massage_cert") else "未填/未具備"),
            ]
        ),
        Spacer(1, 7),
        p("二、服務能力與偏好", heading_style),
        kv(
            [
                ("服務區域", _join(staff.get("regions"))),
                ("可服務時段", _join(staff.get("slots"))),
                ("照護寶寶類型", _join(staff.get("babies"))),
                ("料理/專長", _join(staff.get("skills"))),
                ("交通方式", _join(staff.get("transport"))),
                ("休假偏好", _join(staff.get("rests"))),
            ]
        ),
        Spacer(1, 7),
        p("三、派案資歷摘要", heading_style),
        p(f"目前系統紀錄已排班工作日數：{int(staff.get('scheduled_days') or 0)} 日。"),
        p(f"目前系統紀錄已關聯案件數：{int(staff.get('assigned_cases') or 0)} 件。"),
        p(f"可照護寶寶數上限：{staff.get('care_babies') or '未填'}。"),
        Spacer(1, 7),
        p("四、內部派案備註", heading_style),
        p("本履歷由目前本機測試資料庫自動生成，供 UI/API 驗收、派案流程展示與文件版型測試使用。正式對外使用前，需由行政人員重新核對本人資料、證照、服務經驗與可派案時段。"),
    ]

    def footer(canvas, doc):
        canvas.saveState()
        canvas.setFont(pdf_font, 8)
        canvas.setFillColor(colors.HexColor("#777777"))
        canvas.drawCentredString(A4[0] / 2, 1.0 * cm, f"新竹市月子工會內部測試文件｜生成日 {date.today().isoformat()}")
        canvas.restoreState()

    doc = SimpleDocTemplate(
        str(path),
        pagesize=A4,
        rightMargin=1.8 * cm,
        leftMargin=1.8 * cm,
        topMargin=1.7 * cm,
        bottomMargin=1.7 * cm,
    )
    doc.build(story, onFirstPage=footer, onLaterPages=footer)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--output-dir", required=True)
    parser.add_argument("--limit", type=int, default=10)
    parser.add_argument("--input-json")
    parser.add_argument("--pdf", action="store_true")
    args = parser.parse_args()

    out_dir = Path(args.output_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    if args.input_json:
        staff_rows = json.loads(Path(args.input_json).read_text(encoding="utf-8"))
    else:
        staff_rows = _fetch_staff(args.limit)
    manifest = []
    for index, staff in enumerate(staff_rows, start=1):
        safe_name = str(staff["name"]).replace("/", "_")
        filename = f"{index:02d}_{safe_name}_月嫂履歷.docx"
        path = out_dir / filename
        _resume_doc(staff).save(path)
        item = {"staff_id": staff["id"], "name": staff["name"], "file": filename}
        if args.pdf:
            pdf_filename = f"{index:02d}_{safe_name}_月嫂履歷.pdf"
            _resume_pdf(staff, out_dir / pdf_filename)
            item["pdf"] = pdf_filename
        manifest.append(item)

    (out_dir / "manifest.json").write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps({"output_dir": str(out_dir), "count": len(manifest), "files": manifest}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
